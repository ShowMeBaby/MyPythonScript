#pip install python_docx
# 官网地址
#https://python-docx.readthedocs.io/en/latest/#
import docx
from docx import Document
# 读取文档
document = Document("test.docx")
# 清除操作,目前发现问题为清除会遗留一行空格
document.paragraphs[1].clear()
# 获取段落
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        print("\r\n内容:", run.text, "\r字体大小:",
               'None'if run.font.size==None else run.font.size.pt, "\r字体颜色", run.font.color.rgb)
# 获取表格
for table in document.tables:
	for row in table.rows:
		for cell in row.cells:
			print(cell.text)
# 另存为
document.save('test2.docx')
